"""Build Smart Dental Desk grant proposal as a real editable .docx file.

Run:    python funding/build_doc.py
Output: funding/SmartDentalDesk-GrantProposal.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Palette ─────────────────────────────────────────────────────────────────
NAVY     = RGBColor(0x0B, 0x2A, 0x4A)
BLUE     = RGBColor(0x1D, 0x4E, 0xD8)
GREEN    = RGBColor(0x0A, 0x8A, 0x5A)
INK      = RGBColor(0x1F, 0x2A, 0x3D)
MUTED    = RGBColor(0x47, 0x55, 0x69)
HEADER_FILL = "0B2A4A"
ALT_FILL    = "F3F6FB"

doc = Document()

# ── Base style ──────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# Heading styles
for level, size, color in [(1, 22, NAVY), (2, 16, BLUE), (3, 13, NAVY)]:
    st = doc.styles[f"Heading {level}"]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color


# ── Helpers ─────────────────────────────────────────────────────────────────
def add_para(text="", *, size=11, bold=False, italic=False, color=INK,
             align=None, space_before=0, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = color
    return p


def add_runs(runs, *, align=None, space_before=0, space_after=4):
    """runs: list of (text, {size, bold, italic, color})"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for text, opts in runs:
        r = p.add_run(text)
        r.font.size = Pt(opts.get("size", 11))
        r.font.bold = opts.get("bold", False)
        r.font.italic = opts.get("italic", False)
        r.font.color.rgb = opts.get("color", INK)
    return p


def add_bullet(text, *, size=11, bold=False, color=INK):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def style_cell(cell, text, *, bold=False, color=INK, size=11, align=None, fill=None):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    if fill:
        shade_cell(cell, fill)


def add_table(headers, rows, *, col_widths=None, header_color=RGBColor(0xFF, 0xFF, 0xFF),
              alt_shade=True, accent_col=None):
    """headers: list[str]; rows: list[list[str]]"""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Light Grid Accent 1"
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in tbl.columns[i].cells:
                cell.width = w
    # Header row
    for i, h in enumerate(headers):
        style_cell(tbl.rows[0].cells[i], h,
                   bold=True, color=header_color, size=10.5, fill=HEADER_FILL)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            fill = ALT_FILL if (alt_shade and ri % 2 == 0) else None
            is_accent = (accent_col is not None and ci == accent_col)
            style_cell(tbl.rows[ri + 1].cells[ci], val,
                       bold=is_accent, color=GREEN if is_accent else INK,
                       size=10.5, fill=fill)
    # Space after table
    add_para(space_after=4)
    return tbl


def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CBD5E1")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ──────────────────────────────────────────────────────────────────────────────
# COVER
# ──────────────────────────────────────────────────────────────────────────────
add_para("SMART DENTAL DESK", size=11, bold=True, color=BLUE, space_after=2)
add_para("Grant Funding Proposal", size=30, bold=True, color=NAVY, space_after=4)
add_para("AI-Enabled Dental Practice Management & Patient Engagement Platform",
         size=15, italic=True, color=MUTED, space_after=14)

add_runs([
    ("Funding stage:  ", {"bold": True, "color": MUTED, "size": 11}),
    ("Pre-seed / Seed", {"size": 11}),
    ("        Audience:  ", {"bold": True, "color": MUTED, "size": 11}),
    ("Canadian healthcare-AI grant program", {"size": 11}),
], space_after=2)

add_runs([
    ("Prepared:  ", {"bold": True, "color": MUTED, "size": 11}),
    ("2026-05-22", {"size": 11}),
    ("        Version:  ", {"bold": True, "color": MUTED, "size": 11}),
    ("1.0", {"size": 11}),
], space_after=2)

add_runs([
    ("Contact:  ", {"bold": True, "color": MUTED, "size": 11}),
    ("pvprasanth474@gmail.com", {"size": 11}),
], space_after=12)

add_para("Placeholders marked [TBD] — fill before submission.",
         italic=True, color=MUTED, size=10, space_after=10)
hr()


# ──────────────────────────────────────────────────────────────────────────────
# 1. Executive Summary
# ──────────────────────────────────────────────────────────────────────────────
doc.add_heading("1. Executive Summary", level=1)
add_para(
    "Smart Dental Desk is a production-ready, AI-enabled dental practice management platform "
    "that automates clinical documentation, reduces no-shows, and surfaces operational insights "
    "— built on a country-agnostic architecture ready for Canadian commercialization."
)
add_para(
    "The platform combines a NestJS multi-tenant backend, a Next.js web application, and a "
    "React Native mobile app for clinicians on the go. It is already deployed with paying "
    "clinics in a pilot market (India) and is structurally ready to onboard Canadian dental "
    "practices via a pluggable insurance/regulatory strategy layer (CDCP and Sun Life Canada "
    "providers already seeded)."
)
add_para(
    "Seven AI features are live in production — including AI-assisted clinical notes, "
    "prescription generation with safety checks, treatment plan synthesis, dental chart "
    "analysis, and a rules-based patient insights engine — all powered by Anthropic Claude "
    "and OpenAI APIs with prompt-cache cost optimization."
)
add_runs([
    ("Funding ask:  ", {"bold": True, "color": NAVY, "size": 11}),
    ("CAD $[TBD] over [TBD] months", {"bold": True, "size": 11}),
    (" to (a) complete Canadian regulatory adaptation (PIPEDA, PHIPA, CDA standards), "
     "(b) execute a 3-clinic pilot in Ontario/British Columbia, and (c) launch commercialization.",
     {"size": 11}),
])


# ──────────────────────────────────────────────────────────────────────────────
# 2. Primary Objectives
# ──────────────────────────────────────────────────────────────────────────────
doc.add_heading("2. Primary Objectives", level=1)
add_para("Aligned with the program's commercialization and productivity criteria:")
add_table(
    ["#", "Objective", "Measurable Target", "Status"],
    [
        ["1", "Reduce dental administrative workload",
         "40% reduction vs. baseline", "MVP shipped, pilot data pending"],
        ["2", "Reduce appointment no-shows",
         "30% reduction", "AI Insights + WhatsApp reminders live"],
        ["3", "Improve clinic operational efficiency",
         "25% improvement", "Dashboards + automation cron in production"],
        ["4", "Automate clinical documentation",
         "Voice + AI SOAP notes", "5 of 7 modules live; voice in roadmap"],
        ["5", "Deploy AI patient-engagement systems",
         "Campaigns + WhatsApp + insights", "Live in production"],
        ["6", "Commercialize scalable Canadian healthcare AI",
         "First Canadian clinic by [Q+TBD]", "Architecture ready; pilot pending"],
    ],
    col_widths=[Inches(0.4), Inches(2.6), Inches(2.0), Inches(2.0)],
    accent_col=2,
)


# ──────────────────────────────────────────────────────────────────────────────
# 3. Problem Statement
# ──────────────────────────────────────────────────────────────────────────────
doc.add_heading("3. Problem Statement", level=1)
add_para("Canadian dental practices face four compounding pressures:")
problems = [
    ("Administrative overload.",
     " Front-desk staff spend 30–45% of their shift on appointment management, recall calls, "
     "insurance pre-authorization paperwork, and post-visit follow-up — work that is "
     "repetitive and rules-based."),
    ("No-show economics.",
     " Industry estimates put no-show cost at CAD $150–$250 per missed slot. "
     "A typical 3-operatory clinic loses CAD $[TBD]/year to no-shows alone."),
    ("Clinical documentation burden.",
     " Dentists spend an average of 30 minutes per day re-typing the same SOAP-note phrasings "
     "and prescriptions; mid-career clinicians cite documentation as a top-3 burnout driver."),
    ("Fragmented software stack.",
     " Most existing PMS tools (Dentrix, ABELDent, ClearDent) were architected pre-2010, "
     "run on local servers, lack first-class AI, and charge enterprise prices that lock out "
     "independent practices."),
]
for i, (head, body) in enumerate(problems, 1):
    add_runs([
        (f"{i}. {head}", {"bold": True, "color": NAVY, "size": 11}),
        (body, {"size": 11}),
    ], space_after=4)

add_para(
    "The opportunity: a cloud-native, AI-first platform delivered as a single integrated "
    "product — patient management + clinical + billing + insurance + AI + patient engagement "
    "— at a price point accessible to single-chair clinics.",
    space_before=6,
)


# ──────────────────────────────────────────────────────────────────────────────
# 4. Solution Overview
# ──────────────────────────────────────────────────────────────────────────────
doc.add_heading("4. Solution Overview", level=1)
add_para("Smart Dental Desk is a fully integrated SaaS platform delivered across three surfaces:")

doc.add_heading("4.1  Web application (clinic operations)", level=2)
add_para("Full-featured Next.js application covering:")
for b in [
    "Patient master + medical history + dental chart",
    "Appointment scheduling with conflict detection and recall automation",
    "Clinical visit / consultation flows (SOAP notes, examinations, diagnoses)",
    "Treatment plans with phases, urgency tiers, alternatives, and cost projections",
    "Prescriptions with drug-interaction safety checks",
    "Inventory + medicine stock management",
    "Invoicing, payments, GST/tax, expense tracking",
    "Insurance empanelment, pre-authorization, claims, reimbursement allocation",
    "Multi-channel patient communication: Email, SMS, WhatsApp, in-app",
    "Marketing campaigns + referral tracking + feedback collection",
]:
    add_bullet(b)

doc.add_heading("4.2  Mobile application (clinician-facing)", level=2)
add_para(
    "React Native (Expo) app — patient lookup, appointment view, treatment notes, prescriptions, "
    "billing, and WhatsApp inbox from the dentist's phone."
)

doc.add_heading("4.3  AI feature suite (live in production)", level=2)
ai_features = [
    ("AI Clinical Notes",
     "generates structured SOAP notes from short prompts; populates the same ClinicalVisit model as manual entry."),
    ("AI Treatment Plan",
     "generates multi-phase treatment plans with urgency, alternatives, and cost."),
    ("AI Prescription with safety checks",
     "patient-allergy, age, pregnancy, and drug-interaction aware; inventory-aware (prescribes from clinic stock first)."),
    ("AI Dental Chart Analyzer",
     "Claude/GPT-4 vision: analyzes uploaded dental charts and X-ray images (assistive, not diagnostic)."),
    ("AI Revenue Insights",
     "surfaces revenue anomalies and growth levers from billing data."),
    ("AI Appointment Summary",
     "post-visit summarization."),
    ("AI Campaign Content Generator",
     "drafts multi-language patient campaigns."),
]
for i, (head, body) in enumerate(ai_features, 1):
    add_runs([
        (f"{i}. {head} — ", {"bold": True, "color": NAVY, "size": 11}),
        (body, {"size": 11}),
    ], space_after=3)

add_para(
    "A nightly batch computes Patient Insight Scores (no-show risk, recall due, churn risk, "
    "conversion likelihood) using rules-based scoring in PostgreSQL — zero AI cost on page load, "
    "with LLM reserved for on-click \"explain\" actions.",
    space_before=6,
)

doc.add_heading("4.4  Integrations live today", level=2)
for b in [
    "Razorpay (will be paired with Stripe / Moneris for Canada)",
    "WhatsApp Cloud API (Meta-approved, embedded signup in progress)",
    "Email (transactional + campaign)",
    "File uploads (S3-compatible)",
    "OpenAI + Anthropic Claude",
]:
    add_bullet(b)


# ──────────────────────────────────────────────────────────────────────────────
# 5. AI Relevance
# ──────────────────────────────────────────────────────────────────────────────
doc.add_heading("5. AI Relevance — Real AI, Not Basic Automation", level=1)
add_para("The program asks for \"real AI, not basic automation.\" Our AI surface:")
add_table(
    ["Capability", "Underlying technique", "Why this is real AI"],
    [
        ["Clinical note generation",
         "LLM (Claude + GPT-4) with structured JSON schemas",
         "Context-aware generation across patient history, allergies, complaints — not template substitution"],
        ["Treatment plan synthesis",
         "LLM chain-of-thought with phase / urgency / alternative reasoning",
         "Generates multi-row TreatmentPlan + Item records with internal consistency"],
        ["Prescription safety",
         "LLM + structured inventory lookup + drug-interaction rules",
         "Combines symbolic medical knowledge with retrieval-augmented generation"],
        ["Dental chart vision",
         "LLM vision (Claude Vision, GPT-4V)",
         "Image understanding on radiographs and intraoral photos"],
        ["Voice-to-consultation (roadmap)",
         "OpenAI Whisper + LLM extraction",
         "Speech recognition → structured EHR fields"],
        ["Patient insights scoring",
         "Rules-based scoring in DB; LLM for explanations",
         "Hybrid: deterministic scoring for cost control, LLM for reasoning"],
        ["Cost optimization",
         "Anthropic prompt caching (90% reduction)",
         "Production-grade engineering, not a demo"],
    ],
    col_widths=[Inches(1.8), Inches(2.6), Inches(2.6)],
)
add_runs([
    ("Deliberate exclusion of fake AI.  ", {"bold": True, "color": NAVY, "size": 11}),
    ("We do not claim ML-trained no-show predictors, inventory forecasters, or revenue anomaly "
     "classifiers — those require training infrastructure and data scale we do not yet have. "
     "Where these features appear in the product, they are honestly labeled as rules-based "
     "scoring or LLM-explained heuristics.", {"size": 11}),
])


# ──────────────────────────────────────────────────────────────────────────────
# 6. Productivity & Operational Impact
# ──────────────────────────────────────────────────────────────────────────────
doc.add_heading("6. Productivity & Operational Impact", level=1)
add_table(
    ["Metric", "Baseline (typical Canadian solo practice)",
     "With Smart Dental Desk (target)", "Mechanism"],
    [
        ["Front-desk admin hours/day", "4–5 hrs",
         "2.5–3 hrs (40% reduction)",
         "Recall automation, WhatsApp/SMS reminders, AI-generated communication"],
        ["No-show rate", "12–18%", "8–12% (30% reduction)",
         "Multi-channel reminders + no-show risk scoring + waitlist auto-fill"],
        ["Documentation time per patient", "8–10 min", "3–5 min (50% reduction)",
         "AI clinical notes + voice (roadmap)"],
        ["Prescription error rate", "Industry baseline", "Safety-checked at write",
         "Allergy / interaction / age guardrails"],
        ["Insurance claim turnaround", "14–21 days", "[TBD pilot data]",
         "Pre-auth workflow + claim tracking + reimbursement allocation"],
        ["Patient lifetime value", "Baseline", "+15–25% target",
         "Recall + churn-risk surfacing + conversion of pending plans"],
    ],
    col_widths=[Inches(1.7), Inches(1.7), Inches(1.9), Inches(1.7)],
    accent_col=2,
)
add_para("Targets above will be validated in the Canadian pilot (Section 9).",
         italic=True, color=MUTED)


# ──────────────────────────────────────────────────────────────────────────────
# 7. Technical Readiness
# ──────────────────────────────────────────────────────────────────────────────
doc.add_heading("7. Technical Readiness — MVP, Pilot, and Feasibility", level=1)

doc.add_heading("7.1  Current development phases (completed)", level=2)
add_table(
    ["Phase", "Scope", "Status"],
    [
        ["Phase 1", "NestJS + Prisma + PostgreSQL backend; multi-tenant; RBAC; subscriptions; AI quota", "100% (365+ tests)"],
        ["Phase 2", "Next.js web app — full CRUD, dental chart, invoices, reports, AI features", "100%"],
        ["Phase 2.5", "Frontend hardening — exports, virtual scrolling, Vitest + Playwright", "100%"],
        ["Phase 3", "Email, in-app, campaigns, automation crons, feedback, WhatsApp", "~85%"],
        ["Phase 4", "Docker, CI/CD, Helmet, CSRF, backup, payments, landing page", "100%"],
        ["Phase 5", "7 AI features in production, feature-guarded with quota enforcement", "100%"],
        ["Phase 6", "React Native (Expo) mobile app", "Complete"],
        ["Phase 7", "Insurance & EHS module — 11 Prisma models, country-strategy pattern", "Backend complete; frontend in progress"],
    ],
    col_widths=[Inches(0.9), Inches(4.5), Inches(1.6)],
    accent_col=2,
)

doc.add_heading("7.2  Architecture highlights", level=2)
arch = [
    ("Country-agnostic insurance / regulatory layer.",
     " A CountryStrategy pattern in the insurance module means Canada plugs in as a new "
     "strategy class without schema changes. Canadian providers (CDCP, Sun Life Canada) "
     "are already seeded."),
    ("Multi-tenant by design.",
     " Every entity is scoped by clinic_id with row-level enforcement; supports single "
     "clinics through multi-branch chains."),
    ("Feature flags + plan-gated AI quotas.",
     " Each AI feature is gated by both subscription plan and per-clinic quota; no surprise bills."),
    ("Production hardening.",
     " Helmet, CSRF, rate limiting, automated DB backup, structured logging, observability hooks."),
]
for head, body in arch:
    add_runs([
        ("•  ", {"size": 11, "color": BLUE, "bold": True}),
        (head, {"bold": True, "color": NAVY, "size": 11}),
        (body, {"size": 11}),
    ], space_after=3)

doc.add_heading("7.3  Security & compliance posture", level=2)
add_para(
    "Current state is production-grade for India. For Canadian deployment we will complete:"
)
for b in [
    "PIPEDA & PHIPA alignment (data residency, consent, breach notification)",
    "Canadian-region cloud hosting (AWS Canada Central or Azure Canada Central)",
    "SOC 2 Type I (Year 1) → Type II (Year 2)",
    "Penetration testing + remediation of open code-review findings (15 items tracked internally, all non-blocking)",
]:
    add_bullet(b)


# ──────────────────────────────────────────────────────────────────────────────
# 8. Collaboration & Consortium
# ──────────────────────────────────────────────────────────────────────────────
doc.add_heading("8. Collaboration & Consortium Plan", level=1)
add_para("The program rewards consortium structure. Proposed partners:")
add_table(
    ["Partner type", "Role", "Status"],
    [
        ["Academic — Canadian dental school (Schulich / U of T / UBC)",
         "Clinical validation, research publication, resident training program",
         "[Outreach planned]"],
        ["Clinical — Multi-location DSO or dentist association",
         "Pilot deployment, real-world data, advisory board",
         "[Identifying 3 candidates]"],
        ["Technology — Canadian cloud + AI infra partner",
         "Hosting, data residency, MLOps support",
         "[AWS / Azure CSP engagement]"],
        ["Regulatory — Privacy/compliance consultancy",
         "PIPEDA / PHIPA audit + ongoing compliance",
         "[Quote in progress]"],
        ["Advisor — Practicing Canadian dentist + industry vet",
         "Product-market-fit oversight, intro pipeline",
         "[1 identified, 2 in pipeline]"],
    ],
    col_widths=[Inches(2.5), Inches(2.8), Inches(1.7)],
)
add_para("Letters of intent / collaboration MOUs to be appended in Annex A.",
         italic=True, color=MUTED)


# ──────────────────────────────────────────────────────────────────────────────
# 9. Commercialization Plan
# ──────────────────────────────────────────────────────────────────────────────
doc.add_heading("9. Commercialization Plan", level=1)

doc.add_heading("9.1  Pricing (initial Canadian model)", level=2)
add_para("Mirrors the validated 3-tier structure already running in India, repriced for Canada:")
add_table(
    ["Plan", "India price (validated)", "Canada price (proposed)", "Positioning"],
    [
        ["Free", "₹0", "CAD $0", "Try-before-buy; 1 doctor cap"],
        ["Standard", "₹699/mo (~CAD $11)", "CAD $49–69/mo [TBD]", "Solo/duo practice"],
        ["Growth", "₹1,299/mo (~CAD $21)", "CAD $99–149/mo [TBD]",
         "Multi-op + automation + own WhatsApp Business"],
        ["Enterprise (sales-led)", "Custom", "Custom", "DSO / multi-branch chains"],
    ],
    col_widths=[Inches(1.6), Inches(1.9), Inches(2.0), Inches(1.7)],
    accent_col=2,
)
add_para(
    "Pricing rationale: existing Canadian competitors charge CAD $200–500+/op/month. "
    "Our cloud-native cost structure plus AI quota model allows aggressive entry pricing.",
    italic=True, color=MUTED,
)

doc.add_heading("9.2  Go-to-market sequencing", level=2)
add_runs([("Year 1 (post-funding):", {"bold": True, "color": NAVY, "size": 11})])
for b in [
    "Q1: Canadian regulatory adaptation + Canadian cloud deployment",
    "Q2: 3-clinic pilot (Ontario + BC), KPI measurement",
    "Q3: Public launch — dentist-association partnerships, content + SEO, conferences (ODA, BCDA, Pacific Dental Conference)",
    "Q4: First 50 paying Canadian clinics",
]:
    add_bullet(b)
add_runs([("Year 2:", {"bold": True, "color": NAVY, "size": 11})], space_before=4)
for b in [
    "250 Canadian clinics",
    "DSO partnership pilots",
    "US adjacent-market exploration (US strategies for Delta Dental, MetLife already seeded)",
]:
    add_bullet(b)

doc.add_heading("9.3  Revenue model", level=2)
for b in [
    "SaaS subscription (primary)",
    "WhatsApp / SMS / AI usage overages (Standard + Growth)",
    "Enterprise / DSO contracts (custom)",
    "Future: marketplace fees (lab partners, imaging), insurance claim service fees",
]:
    add_bullet(b)

doc.add_heading("9.4  Unit economics (target)", level=2)
for b in [
    "CAC: CAD $[TBD] (founder-led + content/SEO + association partnerships)",
    "ACV: CAD $[TBD] (blended)",
    "Gross margin: 80%+ (cloud SaaS with capped AI quota)",
    "Payback period: [TBD months]",
    "LTV / CAC target: 4x by Year 2",
]:
    add_bullet(b)


# ──────────────────────────────────────────────────────────────────────────────
# 10. Team
# ──────────────────────────────────────────────────────────────────────────────
doc.add_heading("10. Team", level=1)
add_table(
    ["Role", "Name", "Background"],
    [
        ["Founder / CEO", "[Name]", "[Background]"],
        ["Co-founder / Tech lead", "Prasanth V",
         "Full-stack engineer; built the entire platform; Claude-API caching specialist"],
        ["Clinical advisor", "[Name]", "[Practicing dentist, X years]"],
        ["Canadian market lead", "[To recruit post-funding]", "—"],
        ["Compliance advisor", "[To engage]", "PIPEDA / PHIPA specialist"],
    ],
    col_widths=[Inches(2.0), Inches(2.2), Inches(2.8)],
)
add_para(
    "Hiring plan post-funding: Canadian customer success lead, second backend engineer, "
    "sales / partnerships lead. Detailed org chart in Annex B.",
    italic=True, color=MUTED,
)


# ──────────────────────────────────────────────────────────────────────────────
# 11. Funding Ask & Use of Funds
# ──────────────────────────────────────────────────────────────────────────────
doc.add_heading("11. Funding Ask & Use of Funds", level=1)
add_runs([
    ("Total ask:  ", {"bold": True, "color": NAVY, "size": 12}),
    ("CAD $[TBD] over [TBD] months", {"bold": True, "size": 12, "color": GREEN}),
], space_after=8)
add_para("Indicative allocation (refine to grant program rules):")
add_table(
    ["Bucket", "%", "Notes"],
    [
        ["Engineering (Canadian features, compliance, hardening)", "35%", "2 FTE + contractors"],
        ["Pilot operations + clinical validation", "20%", "3-clinic pilot, data collection, study publication"],
        ["Go-to-market (marketing, content, conferences, sales)", "20%", "Year 1 acquisition engine"],
        ["Compliance + security (SOC 2, PIPEDA audit, pen test)", "10%", "Independent assessors"],
        ["Infrastructure (Canadian cloud, observability)", "5%", "AWS / Azure Canada"],
        ["Working capital + contingency", "10%", "—"],
    ],
    col_widths=[Inches(3.5), Inches(0.8), Inches(2.7)],
    accent_col=1,
)


# ──────────────────────────────────────────────────────────────────────────────
# 12. Risks & Mitigations
# ──────────────────────────────────────────────────────────────────────────────
doc.add_heading("12. Risks & Mitigations", level=1)
add_table(
    ["Risk", "Likelihood", "Impact", "Mitigation"],
    [
        ["Slow regulatory adaptation", "Med", "High",
         "Engage PIPEDA/PHIPA consultant from month 1; pilot under research-collab agreement"],
        ["Canadian market entry slower than projected", "Med", "Med",
         "India revenue cushion + lean ops; partnership-led entry (associations, DSOs)"],
        ["AI cost escalation", "Low", "Med",
         "Prompt caching live (90% reduction); per-plan quotas; deterministic rules for hot paths"],
        ["Competition from incumbents adding AI", "High", "Med",
         "Cloud-native + integrated suite vs. their bolt-on; speed advantage"],
        ["Founder bandwidth", "Med", "High",
         "Hire Canadian market lead + second engineer in month 1"],
        ["Data breach / security incident", "Low", "High",
         "SOC 2, pen test, on-call rotation, breach playbook, cyber insurance"],
    ],
    col_widths=[Inches(2.0), Inches(0.9), Inches(0.8), Inches(3.3)],
)


# ──────────────────────────────────────────────────────────────────────────────
# 13. Milestones & KPIs
# ──────────────────────────────────────────────────────────────────────────────
doc.add_heading("13. Milestones & KPIs (12-month)", level=1)
add_table(
    ["Month", "Milestone"],
    [
        ["M1", "Canadian cloud deployment live; PIPEDA gap analysis complete"],
        ["M2", "3 pilot LOIs signed; consortium partners onboarded"],
        ["M3", "First pilot clinic live in Ontario"],
        ["M4–5", "All 3 pilots running; weekly KPI measurement"],
        ["M6", "Mid-program report: validated reduction in admin time + no-shows"],
        ["M7", "Public Canadian launch"],
        ["M8–10", "First 25 paying clinics"],
        ["M11", "SOC 2 Type I complete"],
        ["M12", "50 paying Canadian clinics; year-2 expansion plan submitted"],
    ],
    col_widths=[Inches(1.0), Inches(6.0)],
    accent_col=0,
)
add_para("KPI dashboard updated monthly; quarterly review with grant program officer.",
         italic=True, color=MUTED)


# ──────────────────────────────────────────────────────────────────────────────
# 14. Annexes
# ──────────────────────────────────────────────────────────────────────────────
doc.add_heading("14. Annexes (to attach before submission)", level=1)
annexes = [
    ("A.", "Letters of intent / collaboration MOUs"),
    ("B.", "Detailed org chart + hiring plan"),
    ("C.", "Technical architecture diagrams (extract from ARCHITECTURE.md)"),
    ("D.", "Security & compliance checklist (extract from SECURITY-CHECKLIST.md)"),
    ("E.", "Product screenshots (web + mobile)"),
    ("F.", "Existing AI feature demo links"),
    ("G.", "Pilot clinic profiles"),
    ("H.", "Detailed financial model (3-year P&L, cap table, runway)"),
]
for tag, desc in annexes:
    add_runs([
        (f"{tag}  ", {"bold": True, "color": BLUE, "size": 11}),
        (desc, {"size": 11}),
    ], space_after=2)

add_para(space_before=12)
hr()
add_para("Prepared by the Smart Dental Desk team.   ·   Follow-up:  pvprasanth474@gmail.com",
         italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, size=10,
         space_before=6)


# ── Save ────────────────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(__file__), "SmartDentalDesk-GrantProposal.docx")
doc.save(out)
print(f"Saved: {out}")
